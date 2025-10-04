from docx import Document
import os

def replace_text_in_docx(docx_file, old_text, new_text):
    doc = Document(docx_file)
    replaced = False

    for paragraph in doc.paragraphs:
        if old_text.lower() in paragraph.text.lower():
            for run in paragraph.runs:
                if old_text.lower() in run.text.lower():
                    run.text = run.text.replace(old_text, new_text)
                    replaced = True
                    break 

            if replaced:
                break 

    if not replaced:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if old_text.lower() in cell.text.lower():
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if old_text.lower() in run.text.lower():
                                    run.text = run.text.replace(old_text, new_text)
                                    replaced = True
                                    break  

                            if replaced:
                                break  

            if replaced:
                break
    if replaced:
        doc.save(docx_file)
        print("Document saved successfully.")
    else:
        print(f"Text '{old_text}' not found in the document.")

print(os.path.dirname(__file__))
